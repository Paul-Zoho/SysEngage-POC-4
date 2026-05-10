"""
pytest fixtures for Source Capture tests.

Provides:
  - fixture_path(name): resolve path to a committed fixture file
  - tmp_fixture(content, suffix): create a temp file with given content
  - simple_requirements_docx: generate simple_requirements.docx via python-docx
  - corrupt_pdf: generate a deliberately corrupt PDF binary
  - very_large_paragraph: generate a >100KB plain-text paragraph
  - project_id / practitioner_id: canonical test identifiers
  - db_cleanup: clean up test entities after each test

Per user preference: .docx and corrupt_pdf generated in conftest (binary-free repo).
Plain-text fixtures (.txt, .md, .xyz) committed as files.
"""

import io
import os
import tempfile
import uuid
from pathlib import Path

import pytest

FIXTURES_DIR = Path(__file__).parent.parent.parent / "fixtures" / "source_capture"

TEST_PROJECT_BASE = "TEST_PROJ"
TEST_PRACTITIONER = "SH_TEST"


@pytest.fixture(scope="session")
def fixture_dir() -> Path:
    return FIXTURES_DIR


def fixture_path(relative: str) -> Path:
    """Return absolute path to a committed fixture file."""
    return FIXTURES_DIR / relative


@pytest.fixture
def project_id() -> str:
    """Unique project ID per test to avoid cross-test contamination."""
    return f"{TEST_PROJECT_BASE}_{uuid.uuid4().hex[:8].upper()}"


@pytest.fixture
def practitioner_id() -> str:
    return TEST_PRACTITIONER


@pytest.fixture
def simple_paragraph_path() -> Path:
    return fixture_path("happy_path/simple_paragraph.txt")


@pytest.fixture
def multi_section_path() -> Path:
    return fixture_path("happy_path/multi_section.md")


@pytest.fixture
def empty_path() -> Path:
    return fixture_path("edge_cases/empty.txt")


@pytest.fixture
def single_short_statement_path() -> Path:
    return fixture_path("edge_cases/single_short_statement.txt")


@pytest.fixture
def abbreviation_path() -> Path:
    return fixture_path("edge_cases/abbreviation_handling.txt")


@pytest.fixture
def simple_requirements_docx(tmp_path: Path) -> Path:
    """
    Generate simple_requirements.docx via python-docx.

    Content: title + 2 heading sections + body paragraphs simulating
    a typical small requirements specification.
    Expected: 2 Segments (H2 headings), 3+ Sources (body paragraphs).
    """
    import docx  # type: ignore[import]

    doc = docx.Document()
    doc.add_heading("System Requirements Specification", level=1)

    doc.add_heading("Functional Requirements", level=2)
    doc.add_paragraph(
        "The system shall accept user input via a web interface. "
        "The interface shall be accessible from standard web browsers."
    )
    doc.add_paragraph(
        "The system shall process requests within five seconds under normal load."
    )

    doc.add_heading("Non-Functional Requirements", level=2)
    doc.add_paragraph(
        "The system shall maintain 99.9% uptime during business hours. "
        "Recovery time objective shall not exceed one hour."
    )

    out_path = tmp_path / "simple_requirements.docx"
    doc.save(str(out_path))
    return out_path


@pytest.fixture
def corrupt_pdf(tmp_path: Path) -> Path:
    """
    Generate a deliberately corrupt PDF.

    Has a valid PDF header so PdfReader attempts to open it, but the
    cross-reference table and content streams are malformed, causing
    partial decode failure (read_completion_status=False).
    """
    corrupt_bytes = (
        b"%PDF-1.4\n"
        + b"1 0 obj\n<< /Type /Catalog >>\nendobj\n"
        + b"CORRUPT_GARBAGE\xff\xfe\x00\x01" * 50
        + b"\n%%EOF\n"
    )
    pdf_path = tmp_path / "corrupt_pdf.pdf"
    pdf_path.write_bytes(corrupt_bytes)
    return pdf_path


@pytest.fixture
def very_large_paragraph(tmp_path: Path) -> Path:
    """
    Generate a >100KB plain-text file with a single paragraph.
    Content is repeated sentences to exceed the 100KB threshold.
    Expected: 1 Source produced; no error.
    """
    sentence = (
        "The system performs analysis of the described product "
        "using the Zachman framework across six perspectives. "
    )
    content = sentence * 1000
    assert len(content.encode("utf-8")) > 100_000, "Generated content must exceed 100KB"
    txt_path = tmp_path / "very_large_paragraph.txt"
    txt_path.write_text(content, encoding="utf-8")
    return txt_path


@pytest.fixture
def unsupported_format_path(tmp_path: Path) -> Path:
    """
    Fixture for unsupported_format.xyz — a binary file that cannot be
    decoded as UTF-8, triggering UnsupportedFormatError.
    Per Implementation Spec §9.2.6.
    """
    # Non-UTF-8 bytes: invalid continuation byte sequence
    binary_content = b"\xff\xfe\x00\x01\x02\x03BINARY_DATA_NOT_UTF8\x80\x81\x82"
    xyz_path = tmp_path / "unsupported_format.xyz"
    xyz_path.write_bytes(binary_content)
    return xyz_path


@pytest.fixture(autouse=True)
def db_cleanup(project_id: str):
    """
    Clean up all test entities after each test.

    Deletes by project_id cascade — removes Sources, SourceAtoms, Segments,
    and AnalysisPasses for the test project.
    """
    yield
    try:
        from core.db import get_session
        from sqlalchemy import delete
        from models.analysis_pass import AnalysisPassModel
        from models.source_atom import SourceAtomModel
        from models.source import SourceModel
        from models.segment import SegmentModel
        from models.project import ProjectModel
        from models.stakeholder import StakeholderModel

        session = get_session()
        try:
            session.execute(
                delete(AnalysisPassModel).where(
                    AnalysisPassModel.project_id == project_id
                )
            )
            session.execute(
                delete(SourceAtomModel).where(
                    SourceAtomModel.project_id == project_id
                )
            )
            session.execute(
                delete(SourceModel).where(SourceModel.project_id == project_id)
            )
            session.execute(
                delete(SegmentModel).where(SegmentModel.project_id == project_id)
            )
            session.execute(
                delete(ProjectModel).where(ProjectModel.project_id == project_id)
            )
            session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()
    except Exception:
        pass
